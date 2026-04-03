#!/usr/bin/env python3
"""Convert RESEARCH_PAPER_DRAFT.md to DOCX with proper tables and embedded figures."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
import re

md = Path("docs/RESEARCH_PAPER_DRAFT.md").read_text()
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Figure mapping - where to insert figures
FIGURE_INSERTIONS = {
    "Summary Dashboard": "analysis/outputs/figures/summary_dashboard.png",
    "ipTM Comparison": "analysis/outputs/charts/iptm_comparison.png",
    "pTM Comparison": "analysis/outputs/charts/ptm_comparison.png",
    "Confidence Intervals": "analysis/outputs/figures/confidence_intervals.png",
    "RMSD Comparison": "analysis/outputs/figures/rmsd_comparison.png",
    "pLDDT Full Chain": "analysis/outputs/figures/plddt_full_chain.png",
    "pLDDT Comparison": "analysis/outputs/figures/plddt_comparison.png",
    "Clinical Targets": "analysis/outputs/figures/clinical_targets.png",
    "PAE Diff Dimer": "analysis/outputs/figures/pae_diff_wt_dimer_vs_compound_dimer.png",
    "WT Mono Structure": "analysis/outputs/figures/structure_job01_wt_mono_fad.png",
    "Compound Dimer Structure": "analysis/outputs/figures/structure_job06_compound_dimer_fad.png",
}

def shade_cells(row, color="1B3A5C"):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        elm = shading.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
        shading.append(elm)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

def add_table(doc, table_lines):
    rows_data = []
    for line in table_lines:
        line = line.strip()
        if line.startswith("|"):
            cells = [c.strip().replace("**", "").replace("*", "") for c in line.split("|")[1:-1]]
            if cells and not all(set(c) <= set("-: ") for c in cells):
                rows_data.append(cells)
    if not rows_data:
        return
    n_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Arial"
    shade_cells(table.rows[0])
    doc.add_paragraph()

def add_figure(doc, path, caption="", width=5.5):
    if Path(path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style.font.size = Pt(9)
            for run in cap.runs:
                run.font.size = Pt(9)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()

lines = md.split("\n")
i = 0
figures_added = False

while i < len(lines):
    line = lines[i].strip()

    if not line or line == "---":
        i += 1
        continue

    # Headings
    if line.startswith("# "):
        doc.add_heading(line[2:].replace("**", ""), level=1)
        i += 1
        continue
    elif line.startswith("## "):
        heading_text = line[3:].replace("**", "")
        doc.add_heading(heading_text, level=2)

        # Insert figures at specific sections
        if "Results" in heading_text and not figures_added:
            pass  # Figures go at end
        i += 1
        continue
    elif line.startswith("### "):
        doc.add_heading(line[4:].replace("**", ""), level=3)
        i += 1
        continue
    elif line.startswith("#### "):
        doc.add_heading(line[5:].replace("**", ""), level=4)
        i += 1
        continue

    # Tables
    if line.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i])
            i += 1
        add_table(doc, table_lines)
        continue

    # Bullet lists
    if line.startswith("- "):
        clean = line[2:].replace("**", "").replace("*", "").replace("`", "")
        doc.add_paragraph(clean, style="List Bullet")
        i += 1
        continue

    # Numbered lists
    if re.match(r"^\d+\.", line):
        clean = re.sub(r"^\d+\.\s*", "", line).replace("**", "").replace("*", "").replace("`", "")
        doc.add_paragraph(clean, style="List Number")
        i += 1
        continue

    # Blockquotes
    if line.startswith("> "):
        clean = line[2:].replace("**", "").replace("*", "").replace("`", "")
        p = doc.add_paragraph(clean)
        p.paragraph_format.left_indent = Inches(0.5)
        i += 1
        continue

    # Regular paragraphs with bold handling
    clean = line.replace("`", "")
    p = doc.add_paragraph()
    parts = clean.split("**")
    for j, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if j % 2 == 1:
                run.bold = True
    i += 1

# Add figures section at the end
doc.add_page_break()
doc.add_heading("Figures", level=1)

figures = [
    ("analysis/outputs/figures/summary_dashboard.png",
     "Figure 1. Summary dashboard of all 64 structural predictions across AlphaFold 3 and Boltz-2 platforms."),
    ("analysis/outputs/figures/clinical_targets.png",
     "Figure 2. Primary experimental follow-up contexts with supporting computational metrics."),
    ("analysis/outputs/charts/iptm_comparison.png",
     "Figure 3. Interface predicted TM-score (ipTM) comparison across variant states."),
    ("analysis/outputs/figures/confidence_intervals.png",
     "Figure 4. Confidence intervals (95% CI) for ipTM and pLDDT at position 429 across all variant states (n=10)."),
    ("analysis/outputs/figures/rmsd_comparison.png",
     "Figure 5. RMSD validation against experimental PDB 6FCX for all variants (all < 2.0 A)."),
    ("analysis/outputs/figures/plddt_full_chain.png",
     "Figure 6. Per-residue pLDDT confidence across all 656 residues for WT, C677T, and compound heterozygous states."),
    ("analysis/outputs/figures/structure_job01_wt_mono_fad.png",
     "Figure 7a. Wild-type monomer with FAD, colored by pLDDT confidence."),
    ("analysis/outputs/figures/structure_job06_compound_dimer_fad.png",
     "Figure 7b. Compound heterozygous dimer with FAD, colored by pLDDT confidence."),
    ("analysis/outputs/figures/pae_diff_wt_dimer_vs_compound_dimer.png",
     "Figure 8. PAE difference map: WT dimer vs compound dimer. Red regions indicate higher predicted error in compound state."),
]

for path, caption in figures:
    if Path(path).exists():
        add_figure(doc, path, caption, width=5.5)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MTHFR Variant Hypothesis Prioritization Program | Igor Mihaljko | DSM.Promo\n")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)
run = p.add_run("ORCID: 0009-0000-1408-1065 | DOI: 10.5281/zenodo.19318627\n")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)
run = p.add_run("GitHub: github.com/DSMPromo/mthfr-target-validation | CC BY-NC-SA 4.0")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

out = Path("docs/MTHFR_Research_Paper.docx")
doc.save(str(out))
print(f"Saved {out} ({out.stat().st_size / 1024:.0f} KB)")
